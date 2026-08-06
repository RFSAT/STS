#!/usr/bin/env python3
"""
Edits the User Guide IN PLACE, in the shooter's own document.

WHY THIS REPLACED A GENERATOR
-----------------------------
The guide used to be rebuilt from scratch by a script that emitted every
paragraph, table and page break itself. That works exactly once. The document
then goes to the author, who lays it out properly in Word — a real template,
real heading styles, the RFSAT logo on the title page, page breaks where a
reader wants them — and the next regeneration throws all of it away and hands
back the script's own idea of a document. It did, twice: the logo and the
styles were lost, and the page breaks were the script's rather than the
author's.

A generator cannot preserve what it does not know about, and it cannot be
taught the whole of a Word template. So it is gone. The document is now the
source, and this script only CHANGES TEXT INSIDE IT: it opens the baseline,
substitutes the strings it is told to substitute, and saves. Styles, images,
headers, footers, numbering, section properties and page breaks are never
touched, because nothing here writes them.

WHAT IT WILL NOT DO
-------------------
Insert paragraphs, insert tables, or move anything. Structural changes are
made by the author in Word, and the result comes back as the new baseline.
This script is for the wording that goes stale between editions — a version
number, a renamed setting, a corrected figure.

RUN AS
------
    python3 edit_guide.py baseline_v1.55.0.docx STS-User-Guide_v1.49.1.docx

Edits are listed in EDITS below, as (old, new) pairs. Every one must match
exactly once across the document, or the script refuses to save: an edit that
silently matches nothing is how a guide comes to describe a setting that was
renamed three releases ago.
"""

import sys
from docx import Document

# (old, new). Applied to paragraph text and to table cell text.
EDITS = [
    ("Version 1.54.0", "Version 1.55.0"),

    # 1.55.0 — the reticle and the lens correction. Appended to paragraphs
    # that already exist; new sections belong in Word.
    ("The guide draws the rings of the selected face over the camera preview, at their "
     "true proportions. Line the printed rings up with the drawn ones.",
     "The guide draws the rings of the selected face over the camera preview, at their "
     "true proportions. Line the printed rings up with the drawn ones. Separately, "
     "Settings carries a reticle — none, a simple crosshair, duplex, mil-dot, MOA grid, "
     "German #4, circle and dot, or an image of your own. Choose None when the camera "
     "looks through a scope: it already shows the scope's own reticle, and a second one "
     "beside it helps nobody. The reticle changes no score; the ring guide is the part "
     "that checks the face."),

    ("If the card is at an angle, the Tilt and rotation controls let you correct it.",
     "A wide lens bows the picture outward, most at the edges and not at all in the "
     "middle, so a ring near the edge measures short. It matters when the card fills "
     "the frame from close range and hardly at all down a range. The app measures it "
     "from the printed rings — they are evenly spaced, so any unevenness across the "
     "picture is the lens — and offers a figure under Lens distortion for you to apply. "
     "For a live stream, measure it once from a photograph taken with the same camera "
     "and enter it in Settings. If the card is at an angle, the Tilt and rotation "
     "controls let you correct it."),




]


def paragraphs(doc):
    """Every paragraph in the body, including the ones inside tables."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_in_paragraph(p, old, new):
    """
    Substitutes within a paragraph while keeping its runs.

    Word splits a sentence into runs wherever anything changes — a spell-check
    mark is enough — so the text being replaced is usually spread across
    several of them. The whole paragraph is joined, the substitution is made
    on that, and the result is written back into the FIRST run with the rest
    emptied. The first run carries the paragraph's formatting, so bold text
    stays bold; what is lost is formatting that varied WITHIN the paragraph,
    which is why a hit is reported and only the paragraphs that need changing
    are ever touched.
    """
    joined = "".join(r.text for r in p.runs)
    if old not in joined:
        return 0
    n = joined.count(old)
    joined = joined.replace(old, new)
    if p.runs:
        p.runs[0].text = joined
        for r in p.runs[1:]:
            r.text = ""
    return n


def main():
    if len(sys.argv) != 3:
        print(__doc__)
        return 2
    src, dst = sys.argv[1], sys.argv[2]
    doc = Document(src)
    problems = []
    for old, new in EDITS:
        hits = sum(replace_in_paragraph(p, old, new) for p in paragraphs(doc))
        print(f"{hits:3d}  {old[:60]!r} -> {new[:60]!r}")
        if hits != 1:
            problems.append(f"{hits} matches for {old!r}, expected exactly 1")
    if problems:
        print("\nNOT SAVED:")
        for x in problems:
            print("  " + x)
        return 1
    doc.save(dst)
    print(f"\nwrote {dst}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
